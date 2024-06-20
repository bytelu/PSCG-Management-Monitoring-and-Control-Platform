import os
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Optional, Tuple

from docx import Document


@dataclass
class RevisionDocs:
    """
    Define una estructura para almacenar datos de revisiones para Auditoría, Intervención y Control Interno.
    """
    auditoria_data: Optional[Dict[str, Dict[str, Dict[str, str]]]] = None
    intervencion_data: Optional[Dict[str, Dict[str, Dict[str, str]]]] = None
    control_interno_data: Optional[Dict[str, Dict[str, Dict[str, str]]]] = None

    def __post_init__(self):
        """
        Inicializa banderas para verificar si hay datos disponibles.
        """
        self.auditoria_band = self.auditoria_data is not None
        self.intervencion_band = self.intervencion_data is not None
        self.control_interno_band = self.control_interno_data is not None


class DataDocs:
    """
    Maneja la lista de datos y proporciona metodos para crear y acceder a diccionarios de datos.
    """

    def __init__(self, data_list: List[str]):
        self.data = {f'P{i + 1:02d}': str(data_list[i]) for i in range(len(data_list))}

    def get(self, key: str) -> Optional[str]:
        """
        Obtiene un valor del diccionario de datos.

        Args:
            key (str): Clave del valor a obtener.

        Returns:
            Optional[str]: Valor correspondiente a la clave o None si no se encuentra.
        """
        return self.data.get(key)

    @classmethod
    def from_list(cls, data_list: List[str]) -> 'DataDocs':
        """
        Crea una instancia de DataDocs a partir de una lista de datos.

        Args:
            data_list (List[str]): Lista de datos a almacenar.

        Returns:
            DataDocs: Instancia de DataDocs creada.
        """
        return cls(data_list)

    @staticmethod
    def create_revision_dict(revision_list: List[Tuple[str, str]], kind: str) \
            -> Optional[Dict[str, Dict[str, Dict[str, str]]]]:
        """
        Crea un diccionario de revisión según el tipo (A, I, C).

        Args:
            revision_list (List[Tuple[str, str]]): Lista de pares de datos (Estatus, Comentario).
            kind (str): Tipo de revisión ("A" para Auditoría, "I" para Intervención, "C" para Control Interno).

        Returns:
            Optional[Dict[str, Dict[str, Dict[str, str]]]]: Diccionario de revisión creado o None si el tipo es inválido.
        """
        if kind in ["A", "I", "C"] and revision_list:
            return {
                kind: {
                    f"{i + 1:02d}": {"E": estado, "C": comentario}
                    for i, (estado, comentario) in enumerate(revision_list)
                }
            }
        return None


def replace_text_in_paragraph(paragraph, regex, data):
    """
    Reemplaza texto dentro de un párrafo basado en una expresión regular y un diccionario de datos.

    Args:
        paragraph (docx.text.paragraph.Paragraph): Párrafo en el que se va a realizar el reemplazo.
        regex (re.Pattern): Patrón de expresión regular compilado para hacer coincidir texto.
        data (dict): Diccionario que mapea el texto a reemplazar (claves) con sus reemplazos (valores).

    Returns:
        None
    """
    for run in paragraph.runs:
        new_text = regex.sub(lambda match: data.get(match.group(0), match.group(0)), run.text)
        if new_text != run.text:
            run.text = new_text


def replace_text_in_table(table, regex, data):
    """
    Reemplaza texto dentro de una tabla basado en una expresión regular y un diccionario de datos.

    Args:
        table (docx.table.Table): Tabla en la que se va a realizar el reemplazo.
        regex (re.Pattern): Patrón de expresión regular compilado para hacer coincidir texto.
        data (dict): Diccionario que mapea el texto a reemplazar (claves) con sus reemplazos (valores).

    Returns:
        None
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, regex, data)


def replace_data(doc, data):
    """
    Reemplaza datos en un documento completo (párrafos y tablas) utilizando un diccionario de datos y expresiones regulares.

    Args:
        doc (docx.Document): Documento en el que se van a realizar los reemplazos.
        data (dict): Diccionario que mapea el texto a reemplazar (claves) con sus reemplazos (valores).

    Returns:
        None
    """
    regex = re.compile('|'.join(re.escape(key) for key in data.keys()))
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, regex, data)
    for table in doc.tables:
        replace_text_in_table(table, regex, data)


def replace_revision(doc: Document, revision: RevisionDocs):
    """
    Maneja el reemplazo de datos específicos de revisión en tablas dentro de un documento.

    Args:
        doc (docx.Document): Documento en el que se van a realizar los reemplazos.
        revision (RevisionDocs): Datos de revisión que contienen los parámetros de auditoría, intervención y control interno.

    Returns:
        None
    """
    for table in doc.tables:
        table_text = table.cell(0, 0).text
        if "AUDITORÍA" in table_text:
            handle_table_replacement(table, revision.auditoria_band, revision.auditoria_data, "A")
        elif "INTERVENCIÓN" in table_text:
            handle_table_replacement(table, revision.intervencion_band, revision.intervencion_data, "I")
        elif "CONTROL INTERNO" in table_text:
            handle_table_replacement(table, revision.control_interno_band, revision.control_interno_data, "C")


def handle_table_replacement(table, band, data, kind):
    """
    Maneja el reemplazo de datos en una tabla según el tipo de revisión y la disponibilidad de datos.

    Args:
        table (docx.table.Table): Tabla en la que se va a realizar el reemplazo.
        band (bool): Indica si hay datos disponibles para el tipo de revisión.
        data (dict): Datos específicos de revisión para un tipo particular.
        kind (str): Tipo de revisión ("A" para Auditoría, "I" para Intervención, "C" para Control Interno).

    Returns:
        None
    """
    if band and data:
        kind_data = data.get(kind, {})
        replacement_dict = {}
        for key, value in kind_data.items():
            replacement_dict[f"E{key}"] = value["E"]
            replacement_dict[f"C{key}"] = value["C"]
        regex = re.compile('|'.join(re.escape(key) for key in replacement_dict.keys()))
        replace_text_in_table(table, regex, replacement_dict)
    else:
        parent = table._element.getparent()
        parent.remove(table._element)


def generate_temp(origen: str, destino: str) -> Optional[str]:
    """
    Genera un archivo temporal a partir de una plantilla.

    Args:
        origen (str): Ruta de la plantilla de archivo.
        destino (str): Ruta del directorio donde se va a generar el archivo temporal.

    Returns:
        Optional[str]: Ruta del archivo temporal generado o None si no se pudo generar.
    """
    temp_destino = os.path.join(destino, "temp.docx")
    try:
        shutil.copy(origen, temp_destino)
        return temp_destino
    except (FileNotFoundError, PermissionError) as e:
        print(f"Error al generar el archivo temporal: {e}")
        return None


def delete_temp(path: str) -> bool:
    """
    Elimina un archivo temporal si existe.

    Args:
        path (str): Ruta del archivo temporal a eliminar.

    Returns:
        bool: True si se eliminó correctamente, False si no se pudo eliminar.
    """
    try:
        Path(path).unlink()
        return True
    except (FileNotFoundError, PermissionError):
        return False


def minuta(data: List[str],
           kind: bool,
           oic: str = '&&&&&&',
           mes: str = '00',
           trimestre: str = '0',
           anyo: str = '0000',
           revision: RevisionDocs = None) -> Optional[str]:
    """
    Genera un archivo docx de tipo minuta de acuerdo con una plantilla predefinida.

    Args:
        data (List[str]): Lista de datos que serán escritos en el docx.
        kind (bool): Tipo de archivo a crear:
                     - True: Papeles de Trabajo
                     - False: Proyectos de Observaciones
        oic (str, optional): Nombre del OIC correspondiente de la minuta. Default es '&&&&&&'.
        mes (str, optional): Mes de la minuta. Default es '00'.
        trimestre (str, optional): Trimestre de la minuta. Default es '0'.
        anyo (str, optional): Año de la minuta. Default es '0000'.
        revision (RevisionDocs, optional): Parámetros de observancia de las actividades de fiscalización.
                                           Obligatorio si kind es False (Proyectos de Observaciones).

    Returns:
        Optional[str]: Ruta donde ha sido guardado el archivo final en caso de éxito, o None si no se pudo generar.
    """
    destino = "C:/Users/angel/Testing/pythonProject1/temp"
    origen = (
        "C:/Users/angel/Testing/pythonProject1/files/minuta_papeles.docx"
        if kind
        else "C:/Users/angel/Testing/pythonProject1/files/minuta_proyectos.docx"
    )

    if not kind and revision is None:
        return None

    temp_path = generate_temp(origen, destino)
    if temp_path:
        doc = Document(temp_path)
        if not kind:
            replace_revision(doc, revision)
        data = DataDocs.from_list(data).data
        replace_data(doc, data)
        delete_temp(temp_path)
        output_path = os.path.join(destino, f"Minuta - {oic} - M{mes}T{trimestre} - {anyo}.docx")
        doc.save(output_path)
        return output_path
    return None


def create_revision(auditoria_values: List[Tuple[str, str]] = None,
                    intervencion_values: List[Tuple[str, str]] = None,
                    control_interno_values: List[Tuple[str, str]] = None):
    """
    Crea una instancia de RevisionDocs con datos de auditoría, intervención y control interno.

    Args:
        auditoria_values (List[Tuple[str, str]], optional): Lista de pares (estatus, comentario) para auditoría.
        intervencion_values (List[Tuple[str, str]], optional): Lista de pares (estatus, comentario) para intervención.
        control_interno_values (List[Tuple[str, str]], optional): Lista de pares (estatus, comentario) para control interno.

    Returns:
        RevisionDocs: Instancia de RevisionDocs con los datos de revisión creados.
    """
    return RevisionDocs(
        auditoria_data=DataDocs.create_revision_dict(auditoria_values, "A"),
        intervencion_data=DataDocs.create_revision_dict(intervencion_values, "I"),
        control_interno_data=DataDocs.create_revision_dict(control_interno_values, "C")
    )


def main():
    """
    Funcion principal de ejemplo de uso
    :return:
    """
    dats = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11,
            12, 13, 14, 15, 16, 17, 18, 19, 20,
            21, 22, 23, 24, 25, 26, 27, 28,
            29, 30, 31, 32, 33, 34, 35]

    revision_list = [
        ["Cumple 1", "Comentario 1"],
        ["Pendiente 2", "Comentario 2"],
        ["Cumple 3", "Comentario 3"],
        ["Pendiente 4", "Comentario 4"],
        ["Cumple 5", "Comentario 5"],
        ["Pendiente 6", "Comentario 6"],
        ["Cumple 7", "Comentario 7"],
        ["Pendiente 8", "Comentario 8"],
        ["Cumple 9", "Comentario 9"],
        ["Pendiente 10", "Comentario 10"],
        ["Cumple 11", "Comentario 11"],
        ["Pendiente 12", "Comentario 12"],
        ["Cumple 13", "Comentario 13"],
        ["Pendiente 14", "Comentario 14"],
        ["Cumple 15", "Comentario 15"],
        ["Pendiente 16", "Comentario 16"],
        ["Cumple 17", "Comentario 17"],
        ["Pendiente 18", "Comentario 18"],
        ["Cumple 19", "Comentario 19"],
        ["Pendiente 20", "Comentario 20"],
        ["Cumple 21", "Comentario 21"],
        ["Cumple 22", "Comentario 22"],
        ["Pendiente 23", "Comentario 23"]
    ]

    revision = create_revision(auditoria_values=revision_list, intervencion_values=revision_list,
                               control_interno_values=revision_list)

    # Papeles de trabajo            ->      kind == True
    # Proyectos de observacion      ->      kind == False
    output = minuta(data=dats, kind=False, revision=revision)
    print(output)


if __name__ == "__main__":
    main()
