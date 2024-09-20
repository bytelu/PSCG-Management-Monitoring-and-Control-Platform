import re
from difflib import SequenceMatcher

import docx
import pandas as pd
from fuzzywuzzy import process


def extract_data_from_df(df, headers, threshold=0.75):
    """
    Extrae datos de un DataFrame basado en headers específicos y un umbral de similitud.

    Args:
        df (pd.DataFrame): DataFrame que contiene los datos.
        headers (list): Lista de encabezados a buscar.
        threshold (float): Umbral de similitud para comparación de cadenas.

    Returns:
        dict: Diccionario con los datos extraídos.
    """
    data = {}

    for header in headers:
        for index, row in df.iterrows():
            for cell in row:
                temp1 = header.lower()
                temp2 = cell.lower()
                if compare_strings(temp1, temp2, threshold):
                    header_index = row.tolist().index(cell)
                    data[header] = row[header_index + 1]
                    break
            else:
                continue
            break
        else:
            data[header] = ''

    return data


def compare_strings(str1, str2, threshold=0.75):
    """
    Compara dos cadenas y devuelve True si la similitud es mayor que el umbral dado.

    Args:
        str1 (str): Primera cadena.
        str2 (str): Segunda cadena.
        threshold (float): Umbral de similitud.

    Returns:
        bool: True si la similitud es mayor que el umbral, False en caso contrario.
    """
    # Verifica que ambas cadenas no estén vacías
    if not str1 or not str2:
        return False

    # Compara las cadenas
    return SequenceMatcher(None, str1, str2).ratio() > threshold or str1 in str2 or str2 in str1


def trim_to_number(trimestre):
    """
    Convierte un trimestre dado en su número correspondiente.

    Args:
        trimestre (str): Cadena que representa un trimestre.

    Returns:
        int or None: Número del trimestre (1 a 4) o None si no se encuentra una coincidencia suficientemente alta.
    """
    trimestres_posibles = ["Primero", "Segundo", "Tercero", "Cuarto"]
    match = process.extractOne(trimestre, trimestres_posibles)

    if match and match[1] >= 80:
        return trimestres_posibles.index(match[0]) + 1
    return None


def extract_trimestre_ejercicio(text):
    """
    Extrae el ejercicio y trimestre de un texto dado.

    Args:
        text (str): Texto del cual extraer el ejercicio y trimestre.

    Returns:
        dict: Diccionario con claves 'Ejercicio' y 'Trimestre'.
    """
    pattern_ejercicio = r'ejercicio (\d{4})'
    pattern_trimestre = r'(\w+)\s+trimestre'

    match_ejercicio = re.search(pattern_ejercicio, text.lower())
    match_trimestre = re.search(pattern_trimestre, text.lower())

    ejercicio = match_ejercicio.group(1) if match_ejercicio else None
    trimestre = trim_to_number(match_trimestre.group(1)) if match_trimestre else None

    return {
        'Ejercicio': ejercicio,
        'Trimestre': trimestre
    }


def extract_inicio_termino(date):
    """
    Extrae las fechas de inicio y término de una cadena de texto.

    Args:
        date (str): Cadena de texto que contiene las fechas.

    Returns:
        dict: Diccionario con claves 'Inicio' y 'Termino'.
    """
    pattern = r"(\d{2}/\d{2}/\d{4})"
    matches = re.findall(pattern, date)

    inicio = None
    termino = None

    if "Término" in date:
        if len(matches) == 1:
            termino = matches[0]
        elif len(matches) > 1:
            inicio = matches[0]
            termino = matches[1]
    else:
        inicio = matches[0] if len(matches) > 0 else None
        termino = matches[1] if len(matches) > 1 else None

    return {'Inicio': inicio, 'Termino': termino}


def extract_number_and_year(number):
    """
    Extrae el número y el año de una cadena de formato específico.

    Args:
        number (str): Cadena que contiene el número y el año en el formato especificado.

    Returns:
        dict: Diccionario con claves 'Numero' y 'Año'.
    """
    pattern = r"([A-Za-z])-([\d]{1,4})/(\d{4})"
    match = re.match(pattern, number)

    if match:
        return {"Numero": match.group(2), "Año": match.group(3)}
    return {"Numero": None, "Año": None}


def extract_tables(dfs):
    """
    Extrae información de las tablas de datos.

    Args:
        dfs (list): Lista de DataFrames que representan las tablas.

    Returns:
        dict: Diccionario con los datos extraídos.
    """
    headers = ['Objetivo', 'Antecedentes', 'Normatividad', 'Importancia y riesgo', 'Alcance y periodo']
    data = {}

    for i, header in enumerate(headers, start=1):
        if i < len(dfs):
            df = dfs[i]
            if not df.empty:
                data[header] = df.iloc[1, 0] if len(df) > 1 else ''
            else:
                data[header] = ''
        else:
            data[header] = ''

    if 'Alcance y periodo' in data:
        data.update(extract_trimestre_ejercicio(data.get('Alcance y periodo', '')))

    return data


def extract_fuerza(df):
    """
    Extrae información sobre la fuerza laboral de un DataFrame.

    Args:
        df (pd.DataFrame): DataFrame que contiene la información de la fuerza laboral.

    Returns:
        dict: Diccionario con los datos extraídos.
    """
    data = {}
    headers = ['Supervisión', 'Responsable', 'Auditores']

    for header in headers:
        data[header] = ''

    data.update(extract_data_from_df(df, headers))

    return data


def extract_pint(path):
    """
    Extrae la información principal del documento especificado.

    Args:
        path (str): Ruta del archivo .docx.

    Returns:
        dict: Diccionario con los datos extraídos.
    """
    try:
        document = docx.Document(path)
        text = document.paragraphs[0].text
        title = 'Planeación considerada para la ejecución de la intervención'
        ratio = SequenceMatcher(None, text, title).ratio()

        if ratio < 0.75:
            return None

        tables = document.tables
        dfs = []

        for table in tables:
            data = []
            for row in table.rows:
                text = [cell.text for cell in row.cells]
                data.append(text)
            df = pd.DataFrame(data)
            dfs.append(df)

        if not dfs:
            return None

        data = extract_header(dfs[0])
        data.update(extract_tables(dfs))
        if len(dfs) > 6:
            data.update(extract_fuerza(dfs[6]))

        return data
    except Exception as e:
        return None


def extract_header(df):
    """
    Extrae el encabezado de la tabla.

    Args:
        df (pd.DataFrame): DataFrame que contiene la tabla con el encabezado.

    Returns:
        dict: Diccionario con los datos extraídos del encabezado.
    """
    headers = [
        'Ente Público', 'Emisor', 'Número', 'Denominación',
        'Clave', 'Tipo de Intervención', 'Área', 'Fecha'
    ]
    data_dict = {}

    if df is not None and len(df) > 4:
        data_dict = extract_data_from_df(df, headers)

    data_dict.update(extract_number_and_year(data_dict.pop('Número', '')))
    data_dict.update(extract_inicio_termino(data_dict.pop('Fecha', '')))
    return data_dict
