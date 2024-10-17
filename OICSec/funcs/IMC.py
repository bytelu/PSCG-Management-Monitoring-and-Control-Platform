import docx
import pandas as pd


def tables_data(tables):
    dfs = []
    for table in tables:
        data = []
        for row in table.rows:
            text = [cell.text for cell in row.cells]
            data.append(text)
        df = pd.DataFrame(data)
        dfs.append(df)
    return dfs


def extract_header(document):
    section = document.sections[0]
    header = section.header
    tables = header.tables
    dfs = tables_data(tables)

    if not dfs:
        return []
    else:
        return dfs


def extract_tables(document):
    tables = document.tables
    dfs = tables_data(tables)

    if not dfs:
        return []
    else:
        return dfs


def extract_oic(df):
    return {"OIC": df.iloc[3,0]}



def extract_datas(df):
    # Variables para almacenar los datos
    data = {
        "Numero": df.iloc[2,0],
        "Clave": df.iloc[2,1],
        "Denominacion": df.iloc[2,3],
        "Año/Trimestre": {
            "Ejecucion": df.iloc[2,5],
            "Conclusion": df.iloc[2,6]
        },
        "Periodo": {
            "Original": df.iloc[5,0],
            "Modificado": df.iloc[5,6]
        },
        "Objeto": {
            "Original": df.iloc[8,1],
            "Modificado": df.iloc[8,6]
        },
        "Alcance": {
            "Original": df.iloc[11,1],
            "Modificado": df.iloc[11,6]
        },
        "Justificacion": df.iloc[12,6]
    }
    return data


def extract_kind(df):
    return {"Tipo": df.iloc[0,1]}


def process_dfs(dfs):
    data = extract_datas(dfs[3])
    data.update(extract_oic(dfs[0]))
    data.update(extract_kind(dfs[2]))
    return data

def read_format_a3(path):
    try:
        document = docx.Document(path)
        dfs = extract_header(document)
        dfs.extend(extract_tables(document))
        return process_dfs(dfs)

    except Exception as e:
        print(e)
        return None
